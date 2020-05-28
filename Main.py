import nltk #библиотека для символьной и статистической обработки естественного языка
from nltk.stem.lancaster import LancasterStemmer #стемер на основе библиотек алгоритма Ланкастера
import numpy as np #библиотека облегчающая работу с двумерными массивами
import tflearn #библиотека поддерживающая нейронную сеть
import tensorflow as tf #библиотека машинного обучения
import random #библиотека генератор рандомного числа
import unicodedata #гарантирует, что символы, которые выглядят одинаково на самом деле идентичны
import sys #Библиотека sys позволяет программисту получать информацию об интерпретаторе Python и операционной системе, работать с вводом и выводом, менять параметры модуля и обрабатывать возникающие ошибки.
import docx # библиотека парсинга текста из Word-файлов
import re # библиотека регулярных выражений, служит для поиска подстрок
import enchant # библиотека для проверки слова на наличие в английском языке, используется для проверки правописания слов и предлагает исправления для слов, написанных с ошибками.
from tkinter import ttk # метод графического интерфейса, оттуда берем виджет ScrollBar и TreeView, widgets (Button, Checkbutton, Entry, Frame, Label, LabelFrame, Menubutton, PanedWindow, Radiobutton, Scale and Scrollbar) to automatically replace the Tk widgets.
from docx.oxml.exceptions import InvalidXmlError # импорт ошибки, для дальнейшего предотвращения из работы программы
import json #библиотека для работы с файлами json, Допустим, у нас есть сложный объект, и мы хотели бы преобразовать его в строку, чтобы отправить по сети или просто вывести для логирования.
import tkinter as tk # библиотека графического интерфейса
from tkinter.filedialog import askdirectory # метод ГИ, спрашиваем директорию у пользователя
import matplotlib.pyplot as plt # библиотека построения графиков
import sys
from PyQt5.QtWidgets import QApplication, QWidget, QInputDialog, QLineEdit, QFileDialog, QPushButton, QListWidget
from PyQt5.QtGui import QIcon, QStandardItemModel, QStandardItem
from PyQt5.QtCore import pyqtSlot
from PyQt5.QtWidgets import QMainWindow, QApplication, QWidget, QAction, QTableWidget, QTableWidgetItem, QVBoxLayout, QTableView
from PyQt5.QtGui import QIcon
from PyQt5.QtCore import pyqtSlot

nltk.download("punkt")

class App(QWidget):
    global tableWidget
    def __init__(self):
        super().__init__()
        self.title = 'Tensorflow CV Recognition'
        self.left = 10
        self.top = 10
        self.width = 640
        self.height = 480
        self.initUI()

    def initUI(self):
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)
        button = QPushButton('Open a folder', self)
        button.setToolTip('This is an example button')
        button.move(100, 70)
        button.clicked.connect(self.on_click)
        self.createTable()

        # Add box layout, add table to box layout and add box layout to widget
        self.layout = QVBoxLayout()
        self.layout.addWidget(button)
        self.layout.addWidget(self.tableWidget)
        self.setLayout(self.layout)

        self.show()
    def openFileNameDialog(self):
        #options = QFileDialog.Options()
        #options |= QFileDialog.DontUseNativeDialog
        #fileName = str(QFileDialog.getExistingDirectory(self, "Select Directory"))
        fileName = QFileDialog.getExistingDirectory(None, 'Select a folder:')
        '''fileName, _ = QFileDialog.getOpenFileName(self, "QFileDialog.getOpenFileName()", "",
                                                  "All Files (*);;Python Files (*.py)", options=options)'''
        if fileName:
            print(fileName)
        #fileName = '/'.join(fileName[0].split("/")[:-1])+"/"
        print(fileName)
        main(fileName)

    def openFileNamesDialog(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        files, _ = QFileDialog.getOpenFileNames(self, "QFileDialog.getOpenFileNames()", "",
                                                "All Files (*);;Python Files (*.py)", options=options)
        if files:
            print(files)

    def createTable(self):
        # Create table
        self.tableWidget = QTableView(self)
        self.model = QStandardItemModel(50,2,self)
        self.model.setHorizontalHeaderLabels(["Document Number", "Decision"])
        self.tableWidget.setModel(self.model)
        self.tableWidget.update()
        #self.model.insertRow(0, [QStandardItem("1"), QStandardItem("2")])

        #self.tableWidget.setColumnCount(2)
        #self.tableWidget.setRowCount(20)
        self.tableWidget.setColumnWidth(0,295)
        self.tableWidget.setColumnWidth(1,295)
        self.tableWidget.move(100, 100)
        self.tableWidget.setUpdatesEnabled(True)
    def insert(self,i,row):
        self.model.insertRow(i-1,[QStandardItem(row[0]),QStandardItem(row[1])])
        self.tableWidget.setModel(self.model)
    @pyqtSlot()
    def on_click(self):
        print('PyQt5 button click')
        self.openFileNameDialog()
#tableWidget = QTableWidget()

# функция определяющая проходит кандидат или нет
def tensorflow_gpu(recognize,time):
    answers = []
    models = ['decision'] # название для сохранения модели
    for i in range(len(models)):
        stemmer = LancasterStemmer() # проводим стеминг для файла из под json
        data = None
        #открываем сам json файл
        with open(models[i]+'.json') as json_data:
            data = json.load(json_data)
        #берем оттуда категории Accepted и Denied
        categories = list(data.keys())
        words = []
        docs = []
        # у каждой категории есть свои предложения в json файле. Они являются оценками, при которых кандидат проходит
        # на работу или же нет. Вот этот текст здесь обрабатывают для приведения в нормальную форму слов
        for each_category in data.keys():
            for each_sentence in data[each_category]:
                w = nltk.word_tokenize(each_sentence)
                words.extend(w)
                docs.append((w, each_category))

        #здесь приводят в исходную форму слова
        words = [stemmer.stem(w.lower()) for w in words]
        words = sorted(list(set(words)))
        training = []
        output_empty = [0] * len(categories)

        #здесь подготавливается начальный тренировочный датасет
        for doc in docs:
            bow = []
            token_words = doc[0]
            token_words = [stemmer.stem(word.lower()) for word in token_words]
            for w in words:
                bow.append(1) if w in token_words else bow.append(0)
            output_row = list(output_empty)
            output_row[categories.index(doc[1])] = 1
            training.append([bow, output_row])
        #генерируют смешание массива
        random.shuffle(training)
        #тренировочный массив кладут в другой массив, так как структура numpy по другому хранит в себе массивы,
        #и видимо обращения к ее элементам производится легче
        training = np.array(training)
        train_x = list(training[:, 0])
        train_y = list(training[:, 1])
        #начинается инициализация и активация нейронной сети
        tf.reset_default_graph()
        net = tflearn.input_data(shape=[None, len(train_x[0])])
        net = tflearn.fully_connected(net, 24)
        net = tflearn.fully_connected(net, 24)
        net = tflearn.fully_connected(net, len(train_y[0]), activation='softmax')
        net = tflearn.regression(net)
        model = tflearn.DNN(net, tensorboard_dir='tflearn_logs')
        #На этой стадии модель обучается по тренировочным датасетам
        if time == 1:
            model.fit(train_x, train_y, n_epoch=25000, batch_size=4, show_metric=True)
            #и сохраняет модель в виде файла в корне проекта
            model.save(models[i]+'_model.tflearn')
            #в дальнейшем сохраненную модель можно загрузить, чтобы не обучать модель постоянно одному и тому же
        else:
            model.load(models[i]+'_model.tflearn')
        #это тоже стемминг - но он для предложений, которых нужно распознать
        def get_tf_record(sentence, words):
            sentence_words = nltk.word_tokenize(sentence)
            sentence_words = [stemmer.stem(word.lower()) for word in sentence_words]
            bow = [0]*len(words)
            for s in sentence_words:
                for i, w in enumerate(words):
                    if w == s:
                        bow[i] = 1
            return(np.array(bow))
        #здесь добавляем в массив результаты
        answers.append(categories[np.argmax(model.predict([get_tf_record(recognize[i],words)]))])
    # и функция их возвращает
    return answers

#функция создания моделей и распознавания CV
def tensorflow(recognize,time):
    answers = []
    models = ['education','work','skills','awards']# название для сохранения модели

    #чтение файла и очистка от лишних слов
    for i in range(len(models)):
        stemmer = LancasterStemmer()# проводим стеминг для файла из под json
        data = None

        # открываем сам json файл
        with open(models[i]+'.json') as json_data:
            data = json.load(json_data)
        #берем оттуда категории Accepted и Denied
        categories = list(data.keys())
        words = []
        docs = []
        # у каждой категории есть свои предложения в json файле. Они являются оценками, при которых кандидат проходит
        # на работу или же нет. Вот этот текст здесь обрабатывают для приведения в нормальную форму слов
        for each_category in data.keys():
            for each_sentence in data[each_category]:
                w = nltk.word_tokenize(each_sentence)
                words.extend(w)
                docs.append((w, each_category))
        # здесь приводят в исходную форму слова
        words = [stemmer.stem(w.lower()) for w in words]
        words = sorted(list(set(words)))
        training = []
        output_empty = [0] * len(categories)

        # здесь подготавливается начальный тренировочный датасет
        for doc in docs:
            bow = []
            token_words = doc[0]
            token_words = [stemmer.stem(word.lower()) for word in token_words]
            for w in words:
                bow.append(1) if w in token_words else bow.append(0)
            output_row = list(output_empty)
            output_row[categories.index(doc[1])] = 1
            training.append([bow, output_row])
        # генерируют смешание массива
        random.shuffle(training)
        # тренировочный массив кладут в другой массив, так как структура numpy по другому хранит в себе массивы,
        # и видимо обращения к ее элементам производится легче
        #подготовка тренировочной модели
        training = np.array(training)
        train_x = list(training[:, 0])
        train_y = list(training[:, 1])
        # начинается инициализация и активация нейронной сети
        tf.reset_default_graph()
        net = tflearn.input_data(shape=[None, len(train_x[0])])
        net = tflearn.fully_connected(net, 24)
        net = tflearn.fully_connected(net, 24)
        net = tflearn.fully_connected(net, len(train_y[0]), activation='softmax')
        net = tflearn.regression(net)
        model = tflearn.DNN(net, tensorboard_dir='tflearn_logs')

        # На этой стадии модель обучается по тренировочным датасетам
        if time == 1:
            model.fit(train_x, train_y, n_epoch=25000, batch_size=4, show_metric=True) #чем больше epoch, тем лучше, но дольше. 25k 3 мин
            # и сохраняет модель в виде файла в корне проекта
            model.save(models[i]+'_model.tflearn')
        else:
            # в дальнейшем сохраненную модель можно загрузить, чтобы не обучать модель постоянно одному и тому же
            model.load(models[i]+'_model.tflearn')

        ##это тоже стемминг - но он для предложений, которых нужно распознать
        def get_tf_record(sentence, words):
            sentence_words = nltk.word_tokenize(sentence)
            sentence_words = [stemmer.stem(word.lower()) for word in sentence_words]
            bow = [0]*len(words)
            for s in sentence_words:
                for i, w in enumerate(words):
                    if w == s:
                        bow[i] = 1
            return(np.array(bow))

        # здесь добавляем в массив результаты
        answers.append(categories[np.argmax(model.predict([get_tf_record(recognize[i],words)]))])
    # и функция их возвращает
    return answers

#обработка текста из cv
def text_extracting(directory):
    global ex
    d = enchant.Dict("en_US")#библиотека проверки английских слов
    total_massiv = dict() #массив для сохранения всех найденных сегментов

    #get_segments - функция поиска сегментов
    def get_segments(filename, doc_name):

        # этой функции вы можете тоже обратить внимание
        # по началу я парсил просто данные, потом смотрю чего-то не хватает, оказалось данные из таблиц не берутся.
        # потом вручную написал функцию парсинга с таблиц
        def show_tables(docx):
            keys = None
            #из всех имеющихся таблиц в документе
            for table in docx.tables:
                data = []
                #проходим по каждой строке
                for i, row in enumerate(table.rows):
                    text = (cell.text for cell in row.cells)
                    if i == 0:
                        keys = tuple(text)
                        continue
                    #и сохраняем данные в виде такой структуры: {"Имя" - "Адиль"}
                    row_data = dict(zip(keys, text))
                    #добавляем все в массив
                    data.append(row_data)
                #здесь я обрабатываю одну ошибку, вам о ней необязательно знать, да и объяснять долго надо, так как сам долго врубался))
                # говоря проще, при парсинге таблиц некоторые данные неправильно ложатся в структуру, пришло исправлять
                if data:
                    variable = False
                    for i in data:
                        for y in i:
                            if y is "Name":
                                variable = True
                    if variable:
                        for i in range(len(data)):
                            v = 1
                            a, b = "", ""
                            for y in data[i]:
                                if v == 1:
                                    a = data[i][y].replace("\n", " ").replace("\t", " ")
                                    v += 1
                                else:
                                    b = y.replace("\n", " ").replace("\t", " ")
                            segment[a] = b
                    else:
                        for i in range(len(data)):
                            for y in data[i]:
                                segment[y.replace("\n", " ").replace("\t", " ")] = data[i][y].replace("\n",
                                                                                                      " ").replace("\t",
                                                                                                                   " ")
                #сверху сохраняем все данные из таблицы в массив для сохранения сегментов, но это не тот массив который
                #был инициализован раньше
        #функция проверки текст на наличие в английском языке
        def eng(massive):
            for i in massive:
                if i:
                    if d.check(i) is False:
                        return False
            return True

        #вот тот массив в который сохраняли табличные данные, здесь будут хранится данные сегментов
        segment = dict()
        name = ""
        #открываем Word файл
        doc = docx.Document(filename)
        segments_dictionary = dict()
        # если в Ворде есть таблицы вызываем функцию парсинга текста из таблицы
        if doc.tables:
            try:
                show_tables(doc)
            except InvalidXmlError:
                pass
            except IndexError:
                pass
        #проходим по всем параграфам документа
        for para in doc.paragraphs:
            for run in para.runs:
                #слово curriculum vitae убираем - мешает
                if run.text.lower().__contains__('curriculum vitae'):
                    continue
                #если находим текст чисто из букв, в то же время имеющийся в английском языке, в то же время если оно
                #больше трех букв и подчеркнуто либо написано большими буквами - то велика вероятность что это название
                #сегмента - сохраняем его
                if (''.join(re.split(';|,|:| |-', run.text)).isalpha()
                    and eng(re.split('; |, |:| |-', run.text))
                    and len(''.join(re.split(';|,|:| |-', run.text))) > 3) \
                        and (run.font.underline or run.text.isupper()):
                    name = run.text
                    segment[name] = ""
                else:
                    if name is "":
                        continue
                    string = run.text.replace("\t", " ").replace("\n", " ")
                    segment[name] += string
        #создаем словари для каждого из сегментов
        segments_dictionary['personal'] = ""
        segments_dictionary['work'] = ""
        segments_dictionary['awards'] = ""
        segments_dictionary['skills'] = ""
        segments_dictionary['education'] = ""
        segments_dictionary['others'] = ""
        #распределяем все найденные предложения в виде сегментов на конкретные сегменты
        for i in segment:
            #если в предложении есть такие слова как personal, birth, nation, sex, martial - велика вероятность что это
            #personal info
            if (i.lower().__contains__("personal") and not i.lower().__contains__("skill")) \
                    or i.lower().__contains__("birth") \
                    or i.lower().__contains__("nation") \
                    or i.lower().__contains__("sex") \
                    or i.lower().__contains__("martial"):
                print("--------------------PERSONAL INFORMATION-------------------------")
                segments_dictionary['personal'] += segment[i]
            #тоже самое с work и остальными - сохраняем все отструктурированные сегменты в словарь
            elif i.lower().__contains__("work") \
                    or i.lower().__contains__("job") \
                    or i.lower().__contains__("employment") \
                    or i.lower().__contains__("industrial") \
                    or i.lower().__contains__("internship") \
                    or i.lower().__contains__("experience") \
                    or (i.lower().__contains__("career") and not i.lower().__contains__("objective")):
                print("----------------------WORK EXPERIENCE----------------------------")
                segments_dictionary['work'] += segment[i]
            elif i.lower().__contains__("award") \
                    or i.lower().__contains__("achievement") \
                    or i.lower().__contains__("training"):
                print("--------------------AWARDS AND ACHIEVEMENTS----------------------")
                segments_dictionary['awards'] += segment[i]
            elif i.lower().__contains__("skill") \
                    or i.lower().__contains__("technique"):
                print("---------------------------SKILLS--------------------------------")
                segments_dictionary['skills'] += segment[i]
            elif i.lower().__contains__("education") \
                    or i.lower().__contains__("discipline") \
                    or i.lower().__contains__("learning") \
                    or i.lower().__contains__("study") \
                    or i.lower().__contains__("science") \
                    or i.lower().__contains__("academic") \
                    or i.lower().__contains__("diploma") \
                    or i.lower().__contains__("knowledge") \
                    or i.lower().__contains__("certificat") \
                    or i.lower().__contains__("institution") \
                    or i.lower().__contains__("university") \
                    or i.lower().__contains__("course"):
                print("--------------------------EDUCATION------------------------------")
                segments_dictionary['education'] += segment[i]
            else:
                print("----------------------------OTHER--------------------------------")
                segments_dictionary['others'] += segment[i]
            print(i, "\t", segment[i])
            # print()
            #массив, который мы возвращаем в конце функции
            total_massiv[int(doc_name)] = segments_dictionary
    #это цикл для прохождения по всем word файлам
    for i in range(1, 20):
        try:
            #вызываем фунцию поиска сегментов
            get_segments(
                directory+"/" + str(
                    i) + ".docx", i)
            connect = []
            #берем все сегменты кроме personal и others для обучение нейронной сети
            for y in total_massiv[i]:
                if y != 'personal' or y != 'others':
                    connect.append(total_massiv[i][y])
            #обучаем нейронку - берем результаты сегментов
        except KeyError:
            ex.model.insertRow(i, [QStandardItem(str(i+1)), QStandardItem(str("Not Found"))])
            ex.update()
            continue
        answers = tensorflow(connect,i+1)
        one, two = str(i+1),tensorflow_gpu(answers,i+1)[0]
        print(one,"\t",two)
        #сохраняем в список в графическом приложении результаты Accepted and Denied
        '''ex.tableWidget.setItem(1,0,QTableWidgetItem(one))
        ex.tableWidget.setItem(1,1,QTableWidgetItem(two))
        ex.tableWidget.update()'''
        #ex.insert(i,[one,two])
        ex.model.insertRow(i,[QStandardItem(str(one)), QStandardItem(str(two))])
        ex.update()
        #(ex.tableWidget.model()).appendRow([QStandardItem(one),QStandardItem(two)])
        #ex.model.insertRow(i, [QStandardItem("1"), QStandardItem("2")])
        #ex.tableWidget.update()

    #возвращаем массив с сегментами
    return total_massiv

#главная функция для вызова всех функции
def main(filename):
    #спрашиваем у пользователя какую директорию он выбирает
    #сохраняем в graph данные сегментов, чтобы выразить их в графике
    graphic = text_extracting(filename)
    gr = dict()
    gr["personal"] = 0
    gr["work"] = 0
    gr["skills"] = 0
    gr["awards"] = 0
    gr["education"] = 0
    #ищем сколько сегментов у каждого резюме

    for i in range(1,20):
        try:
            if len(graphic[i]["personal"])!=0:
                gr["personal"] += 1
            elif len(graphic[i]["work"])!=0:
                gr["work"] += 1
            elif len(graphic[i]["skills"])!=0:
                gr["skills"] += 1
            elif len(graphic[i]["awards"])!=0:
                gr["awards"] += 1
            elif len(graphic[i]["education"])!=0:
                gr["education"] += 1
        except KeyError:
            continue
    print(gr)
    #выводим в виде графа
    plt.bar(range(len(gr)), list(gr.values()), align='center')
    plt.xticks(range(len(gr)), list(gr.keys()))
    plt.show()

app = QApplication(sys.argv)
ex = App()
sys.exit(app.exec_())